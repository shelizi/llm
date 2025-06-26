"""Web UI for RAG file 管理與索引建立。
此模組在匯入後會把路由掛載到現有的 `rag_api.app`。
"""

from __future__ import annotations

from pathlib import Path
import logging
import shutil
from typing import Dict, List

import chromadb
import json
import docx
from pypdf import PdfReader
from fastapi import BackgroundTasks, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import HTMLResponse, RedirectResponse
from fastapi.templating import Jinja2Templates

from llama_index.core import Document
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core.text_splitter import TokenTextSplitter
from llama_index.core import StorageContext, VectorStoreIndex
from llama_index.vector_stores.chroma import ChromaVectorStore

import torch

# 取用既有的 FastAPI instance 與設定
import rag_api  # noqa
from rag_api import app, CHROMA_DIR, COLLECTION_NAME, EMBED_MODEL_NAME, DEVICE, index, RAG_EMBEDDING_DIR

# ----------------- 路徑與模板 -----------------
BASE_DIR = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)
TEMPLATES_DIR = BASE_DIR / "templates"
TEMPLATES_DIR.mkdir(exist_ok=True)

templates = Jinja2Templates(directory=str(TEMPLATES_DIR))

# ----------------- 狀態管理 -----------------
STATUS_FILE = BASE_DIR / "status.json"

# 讀取現有狀態，若無則建立空 dict
if STATUS_FILE.exists():
    try:
        StatusDict: Dict[str, str] = json.loads(STATUS_FILE.read_text(encoding="utf-8"))
    except Exception:
        StatusDict: Dict[str, str] = {}
else:
    StatusDict: Dict[str, str] = {}

def save_status() -> None:
    """將 StatusDict 寫入磁碟"""
    try:
        STATUS_FILE.write_text(json.dumps(StatusDict, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as e:
        logging.error("無法寫入狀態檔: %s", e)

# ----------------- 輔助函式 -----------------

def read_file_text(filepath: Path) -> str:
    """讀取檔案文本，包含錯誤處理"""
    suffix = filepath.suffix.lower()
    
    try:
        if suffix == ".txt":
            text = filepath.read_text(encoding="utf-8", errors="ignore")
        elif suffix == ".pdf":
            reader = PdfReader(str(filepath))
            text_parts = []
            total_pages = len(reader.pages)
            logging.info(f"PDF 檔案共 {total_pages} 頁")
            
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())
                        logging.debug(f"PDF 第 {i+1} 頁提取成功，長度: {len(page_text)}")
                    else:
                        logging.warning(f"PDF 第 {i+1} 頁無法提取文本或內容為空")
                except Exception as e:
                    logging.warning(f"PDF 第 {i+1} 頁處理失敗: {e}")
                    continue
            
            if not text_parts:
                raise ValueError("PDF 檔案中沒有可提取的文本內容")
                
            text = "\n\n".join(text_parts)  # 使用雙換行分隔頁面
            logging.info(f"PDF 處理完成，成功提取 {len(text_parts)} 頁內容")
        elif suffix in (".docx", ".doc"):
            document = docx.Document(str(filepath))
            text_parts = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            text = "\n".join(text_parts)
        else:
            raise ValueError(f"不支援的檔案格式: {suffix}")
        
        # 確保文本不為空
        if not text or not text.strip():
            raise ValueError("檔案內容為空或無法提取文本")
            
        # 清理文本
        text = text.strip()
        if len(text) < 10:  # 文本太短可能有問題
            raise ValueError(f"提取的文本太短 ({len(text)} 字符)，可能檔案損壞")
            
        logging.info(f"成功提取文本，長度: {len(text)} 字符")
        return text
        
    except Exception as e:
        logging.error(f"讀取檔案 {filepath.name} 失敗: {e}")
        raise


def process_and_index_file(filename: str, filepath: Path, chunk_size: int, overlap: int):
    """背景任務：處理檔案並寫入向量資料庫"""
    try:
        logging.info(f"開始處理檔案: {filename}")
        StatusDict[filename] = "索引中"
        save_status()
        
        # 驗證參數
        if chunk_size <= 0:
            chunk_size = 512
        if overlap < 0:
            overlap = 50
        if overlap >= chunk_size:
            overlap = chunk_size // 4  # 設為 chunk_size 的 1/4
            
        logging.info(f"使用參數 - chunk_size: {chunk_size}, overlap: {overlap}")
        
        # 讀取文件內容
        StatusDict[filename] = "讀取中"
        save_status()
        text = read_file_text(filepath)
        
        # 創建文檔
        StatusDict[filename] = "建立文檔中"
        save_status()
        documents = [Document(text=text, metadata={"filename": filename, "file_path": str(filepath)})]
        logging.info(f"建立文檔完成，文本長度: {len(text)}")

        # 準備向量存取
        StatusDict[filename] = "準備向量存取中"
        save_status()
        
        # 修復 ChromaDB 客戶端初始化問題
        try:
            # 使用 PersistentClient 而非 HttpClient
            client = chromadb.PersistentClient(path=str(CHROMA_DIR))
            logging.info(f"ChromaDB 客戶端初始化成功，使用路徑: {CHROMA_DIR}")
        except Exception as e:
            logging.error(f"ChromaDB 客戶端初始化失敗: {e}")
            # 嘗試使用備用方法初始化
            client = chromadb.Client()
            logging.info("使用備用方法初始化 ChromaDB 客戶端")
            
        # 創建或獲取集合
        try:
            # 先檢查集合是否存在
            try:
                collection = client.get_collection(name=COLLECTION_NAME)
                logging.info(f"獲取現有集合: {COLLECTION_NAME}")
            except Exception:
                # 如果不存在則創建
                collection = client.create_collection(name=COLLECTION_NAME)
                logging.info(f"創建新集合: {COLLECTION_NAME}")
                
            # 使用集合創建 vector_store
            vector_store = ChromaVectorStore(
                chroma_collection=collection,
                collection_name=COLLECTION_NAME
            )
        except Exception as e:
            logging.error(f"創建向量存儲失敗: {e}")
            # 嘗試直接使用 client 創建
            vector_store = ChromaVectorStore(
                chroma_client=client,
                collection_name=COLLECTION_NAME
            )
        storage_context = StorageContext.from_defaults(vector_store=vector_store)
        
        # 檢查並下載 embedding 模型
        StatusDict[filename] = "檢查模型中"
        save_status()
        
        try:
            from rag_api import check_and_download_embedding_model
            if not check_and_download_embedding_model(EMBED_MODEL_NAME):
                raise RuntimeError(f"無法載入或下載 embedding 模型: {EMBED_MODEL_NAME}")
        except ImportError:
            # 如果無法從 rag_api 導入，則在此處定義檢查函數
            logging.warning("無法從 rag_api 導入檢查函數，使用本地檢查")
            try:
                # 嘗試創建 embedding 模型實例來檢查
                test_embed_model = HuggingFaceEmbedding(model_name=EMBED_MODEL_NAME, device=DEVICE, cache_folder=str(RAG_EMBEDDING_DIR))
            except Exception as e:
                logging.warning(f"⚠️ Embedding 模型載入失敗，嘗試下載到專案目錄: {e}")
                try:
                    from sentence_transformers import SentenceTransformer
                    import os
                    
                    # 設置專案目錄下的模型路徑
                    models_dir = BASE_DIR / "models"
                    embedding_dir = models_dir / "embedding"
                    embedding_dir.mkdir(parents=True, exist_ok=True)
                    
                    # 設置環境變數
                    os.environ['SENTENCE_TRANSFORMERS_HOME'] = str(embedding_dir)
                    
                    # 下載模型到專案目錄
                    logging.info(f"📥 下載模型到專案目錄: {embedding_dir}")
                    model = SentenceTransformer(EMBED_MODEL_NAME, cache_folder=str(embedding_dir))
                    logging.info(f"✅ 成功下載 embedding 模型到專案目錄: {EMBED_MODEL_NAME}")
                except Exception as download_error:
                    raise RuntimeError(f"無法下載 embedding 模型: {download_error}")
        
        embed_model = HuggingFaceEmbedding(model_name=EMBED_MODEL_NAME, device=DEVICE, cache_folder=str(RAG_EMBEDDING_DIR))
        
        # 創建文本分割器，添加錯誤處理
        try:
            text_splitter = TokenTextSplitter(
                chunk_size=chunk_size, 
                chunk_overlap=overlap,
                separator=" "  # 明確指定分隔符
            )
            logging.info("文本分割器創建成功")
        except Exception as e:
            logging.error(f"創建文本分割器失敗: {e}")
            # 使用預設參數重試
            text_splitter = TokenTextSplitter(chunk_size=512, chunk_overlap=50)
            logging.info("使用預設參數重新創建文本分割器")

        # 建立或更新索引
        StatusDict[filename] = "建立索引中"
        save_status()
        
        global index  # 取用 rag_api 的 index 變數
        if index is None:
            logging.info("建立新的索引")
            index = VectorStoreIndex.from_documents(
                documents,
                storage_context=storage_context,
                embed_model=embed_model,
                text_splitter=text_splitter,
                show_progress=True
            )
        else:
            logging.info("插入文檔到現有索引")
            # 已有索引 -> 直接插入文件
            for doc in documents:
                index.insert(doc, text_splitter=text_splitter)

        # 將索引同步到 rag_api 供 /query 使用
        rag_api.index = index  # 確保查詢端點可用最新索引
        StatusDict[filename] = "已索引"
        save_status()
        logging.info("✅ %s 已完成索引", filename)
        
    except Exception as e:
        error_msg = f"錯誤: {str(e)}"
        StatusDict[filename] = error_msg
        save_status()
        logging.error("❌ 索引檔案 %s 失敗: %s", filename, e)
        # 記錄詳細的錯誤信息
        import traceback
        logging.error("詳細錯誤信息:\n%s", traceback.format_exc())


# ----------------- 路由 -----------------

@app.get("/")
async def root():
    """根路由，重定向到 UI"""
    return RedirectResponse("/ui", status_code=302)

@app.get("/ui", response_class=HTMLResponse)
async def ui_root(request: Request):
    try:
        files = sorted([p.name for p in UPLOAD_DIR.iterdir() if p.is_file()])
        statuses = {f: StatusDict.get(f, "未索引") for f in files}
        logging.info(f"UI 頁面載入，找到 {len(files)} 個檔案")
        return templates.TemplateResponse("index.html", {"request": request, "files": files, "statuses": statuses})
    except Exception as e:
        logging.error(f"UI 頁面載入失敗: {e}")
        raise HTTPException(status_code=500, detail=f"UI 載入錯誤: {str(e)}")


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    dest = UPLOAD_DIR / file.filename
    with dest.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    StatusDict[file.filename] = "已上傳"
    save_status()
    return RedirectResponse("/ui", status_code=303)


@app.post("/index")
async def index_files(
    background_tasks: BackgroundTasks,
    filenames: List[str] = Form(...),
    chunk_size: int = Form(512),
    overlap: int = Form(50)
):
    """接受多個檔案名稱，批次排入索引任務"""
    if not filenames:
        raise HTTPException(status_code=400, detail="未提供檔案名稱")

    for filename in filenames:
        filepath = UPLOAD_DIR / filename
        if not filepath.exists():
            logging.warning("檔案 %s 不存在，跳過索引", filename)
            continue
        background_tasks.add_task(process_and_index_file, filename, filepath, chunk_size, overlap)
    return RedirectResponse("/ui", status_code=303)


@app.get("/health")
async def health_check():
    """健康檢查端點"""
    return {
        "status": "ok", 
        "upload_dir": str(UPLOAD_DIR),
        "templates_dir": str(TEMPLATES_DIR),
        "upload_dir_exists": UPLOAD_DIR.exists(),
        "templates_dir_exists": TEMPLATES_DIR.exists(),
        "index_html_exists": (TEMPLATES_DIR / "index.html").exists()
    }

@app.get("/status")
async def status():
    return StatusDict

@app.post("/reset_status")
async def reset_status(filename: str = Form(...)):
    """重置特定檔案的狀態"""
    if filename in StatusDict:
        StatusDict[filename] = "未索引"
        save_status()
        logging.info(f"重置檔案狀態: {filename}")
    return RedirectResponse("/ui", status_code=303)

@app.post("/delete_file")
async def delete_file(filename: str = Form(...)):
    """刪除指定檔案並移除狀態"""
    filepath = UPLOAD_DIR / filename
    try:
        if filepath.exists():
            filepath.unlink()
            logging.info("已刪除檔案: %s", filename)
    except Exception as e:
        logging.error("刪除檔案 %s 失敗: %s", filename, e)

    if filename in StatusDict:
        del StatusDict[filename]
        save_status()

    return RedirectResponse("/ui", status_code=303)

@app.post("/clear_all_status")
async def clear_all_status():
    """清除所有檔案狀態"""
    StatusDict.clear()
    save_status()
    logging.info("清除所有檔案狀態")
    return RedirectResponse("/ui", status_code=303)


# ----------------- 啟動服務器 -----------------
if __name__ == "__main__":
    import uvicorn
    logging.info("🚀 啟動 RAG Web UI 服務器...")
    uvicorn.run(app, host="0.0.0.0", port=8000)
